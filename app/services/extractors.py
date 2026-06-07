from pathlib import Path


def extract_markdown(path: Path, limit: int = 1_500_000) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")[:limit]
    except OSError:
        return ""


def extract_plain_text(path: Path, limit: int = 1_500_000) -> str:
    for encoding in ("utf-8", "gb18030", "gbk"):
        try:
            return path.read_text(encoding=encoding, errors="strict")[:limit]
        except (UnicodeDecodeError, OSError):
            continue
    try:
        return path.read_text(encoding="utf-8", errors="ignore")[:limit]
    except OSError:
        return ""


def _ocr_pdf(path: Path, limit: int = 1_500_000) -> str:
    from app.config import settings

    if not settings.pdf_ocr_enabled:
        return "[PDF text extraction returned empty. OCR is disabled. Set RAG_PDF_OCR_ENABLED=true to enable it.]"
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return f"[PDF OCR unavailable: install pymupdf, pillow and pytesseract. Missing: {exc}]"

    try:
        document = fitz.open(str(path))
        pages: list[str] = []
        max_pages = max(1, settings.pdf_ocr_max_pages)
        for page_index in range(min(len(document), max_pages)):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
            if text:
                pages.append(f"# OCR Page {page_index + 1}\n\n{text}")
            if sum(len(item) for item in pages) >= limit:
                break
        return "\n\n".join(pages)[:limit] or "[PDF OCR completed but no text was recognized.]"
    except Exception as exc:  # noqa: BLE001
        return f"[PDF OCR failed: {exc}]"


def extract_pdf(path: Path, limit: int = 1_500_000) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[PDF extraction unavailable: install pypdf]"

    try:
        reader = PdfReader(str(path))
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"# Page {index}\n\n{text.strip()}")
            if sum(len(item) for item in pages) >= limit:
                break
        extracted = "\n\n".join(pages)[:limit]
        return extracted if extracted.strip() else _ocr_pdf(path, limit=limit)
    except Exception as exc:  # noqa: BLE001
        return f"[PDF extraction failed: {exc}]"


def extract_docx(path: Path, limit: int = 1_500_000) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[DOCX extraction unavailable: install python-docx]"

    try:
        document = Document(str(path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
            if sum(len(item) for item in blocks) >= limit:
                break
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)[:limit]
    except Exception as exc:  # noqa: BLE001
        return f"[DOCX extraction failed: {exc}]"


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return extract_markdown(path)
    if suffix == ".txt":
        return extract_plain_text(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    return ""


def document_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return "markdown"
    if suffix == ".txt":
        return "txt"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    return "unknown"
